"""
Script para gerar documento Word com documentacao completa do Sistema de Demanda e Reabastecimento v5.0
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_cell_shading(cell, color):
    """Define cor de fundo de uma celula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_header(doc, headers, data, header_color='667eea'):
    """Adiciona tabela formatada com cabecalho colorido"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Cabecalho
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    # Dados
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def criar_documento():
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========================================
    # CAPA
    # ========================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    titulo = doc.add_paragraph()
    titulo_run = titulo.add_run("Sistema de Demanda e Reabastecimento")
    titulo_run.bold = True
    titulo_run.font.size = Pt(28)
    titulo_run.font.color.rgb = RGBColor(102, 126, 234)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    versao = doc.add_paragraph()
    versao_run = versao.add_run("Versao 5.0")
    versao_run.font.size = Pt(20)
    versao_run.font.color.rgb = RGBColor(100, 100, 100)
    versao.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitulo = doc.add_paragraph()
    sub_run = subtitulo.add_run("Documentacao Tecnica e Manual do Usuario")
    sub_run.font.size = Pt(16)
    sub_run.italic = True
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    empresa = doc.add_paragraph()
    emp_run = empresa.add_run("Ferreira Costa e Cia Ltda")
    emp_run.font.size = Pt(14)
    emp_run.bold = True
    empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = doc.add_paragraph()
    data_run = data.add_run(f"Janeiro 2026")
    data_run.font.size = Pt(12)
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========================================
    # INDICE
    # ========================================
    doc.add_heading('Indice', level=1)

    indice_items = [
        "1. Introducao",
        "2. Caracteristicas Principais",
        "3. Arquitetura do Sistema",
        "4. Instalacao e Configuracao",
        "5. Modulos do Sistema",
        "   5.1. Previsao de Demanda",
        "   5.2. Pedido ao Fornecedor Integrado",
        "   5.3. Transferencias entre Lojas",
        "   5.4. Simulador de Cenarios",
        "   5.5. Gerenciamento de Eventos",
        "   5.6. KPIs e Metricas",
        "   5.7. Pedido Manual",
        "   5.8. Central de Parametros",
        "6. Politica de Estoque ABC",
        "7. APIs Disponiveis",
        "8. Estrutura do Banco de Dados",
        "9. FAQ - Perguntas Frequentes",
        "10. Changelog",
        "11. Suporte e Contato"
    ]

    for item in indice_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========================================
    # 1. INTRODUCAO
    # ========================================
    doc.add_heading('1. Introducao', level=1)

    intro_text = """O Sistema de Demanda e Reabastecimento v5.0 e uma solucao completa para gestao de estoque multi-loja com Centro de Distribuicao (CD), combinando previsao de demanda Bottom-Up com politica de estoque baseada em curva ABC.

Esta ferramenta foi desenvolvida para otimizar o processo de compras e reabastecimento, utilizando algoritmos estatisticos avancados para prever a demanda futura e calcular automaticamente as quantidades ideais de pedido."""

    doc.add_paragraph(intro_text)

    doc.add_heading('Principais Novidades da Versao 5.0', level=2)

    novidades = [
        "Pedido Multi-Loja com layout hierarquico (Fornecedor > Loja > Itens)",
        "Transferencias inteligentes entre lojas do mesmo grupo regional",
        "Parametros de fornecedor (Lead Time, Ciclo, Faturamento Minimo)",
        "Graceful degradation para novas tabelas",
        "Multi-selecao em filtros (Lojas, Fornecedores, Categorias)",
        "Ordenacao por criticidade de alertas na tabela de previsao detalhada",
        "Coluna de Situacao de Compra (SIT_COMPRA) na previsao de demanda"
    ]

    for item in novidades:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # 2. CARACTERISTICAS PRINCIPAIS
    # ========================================
    doc.add_heading('2. Caracteristicas Principais', level=1)

    doc.add_heading('2.1 Previsao de Demanda V2 (Bottom-Up)', level=2)

    prev_text = """O sistema utiliza uma abordagem Bottom-Up que analisa cada produto individualmente e seleciona o melhor metodo estatistico automaticamente."""
    doc.add_paragraph(prev_text)

    # Tabela de metodos
    doc.add_paragraph()
    doc.add_paragraph("Metodos Estatisticos Disponiveis:").bold = True

    metodos_headers = ['Metodo', 'Descricao', 'Quando e Usado']
    metodos_data = [
        ['SMA', 'Simple Moving Average', 'Demanda estavel, baixa variabilidade'],
        ['WMA', 'Weighted Moving Average', 'Demanda com tendencia recente'],
        ['SES', 'Simple Exponential Smoothing', 'Demanda com ruido moderado'],
        ['Linear Regression', 'Regressao Linear', 'Tendencia clara de crescimento/queda'],
        ['TSB', 'Trigg-Leach Smoothing', 'Demanda intermitente'],
        ['Decomposicao Sazonal', 'Hibrida', 'Padroes sazonais detectados']
    ]
    add_table_with_header(doc, metodos_headers, metodos_data)

    doc.add_paragraph()
    doc.add_heading('Diferenciais do Sistema:', level=3)

    diferenciais = [
        "Selecao automatica do melhor metodo por SKU",
        "Limitadores de tendencia (queda >40% limitada a -40%, alta >50% limitada a +50%)",
        "Sazonalidade anual (12 meses ou 52 semanas)",
        "Deteccao automatica de outliers (IQR + Z-Score)",
        "Saneamento de rupturas: substitui zeros de ruptura por valores estimados"
    ]

    for item in diferenciais:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 Modulos Disponiveis', level=2)

    modulos_headers = ['Modulo', 'Descricao', 'Status']
    modulos_data = [
        ['Previsao de Demanda', 'Forecasting Bottom-Up multi-metodo', 'Ativo'],
        ['Simulador', 'Cenarios what-if', 'Ativo'],
        ['Eventos', 'Calendario promocional', 'Ativo'],
        ['KPIs e Metricas', 'Dashboard de performance', 'Ativo'],
        ['Pedido ao Fornecedor', 'Calculo integrado multi-loja', 'Atualizado v5.0'],
        ['Transferencias entre Lojas', 'Balanceamento automatico de estoque', 'Novo v5.0'],
        ['Parametros Fornecedor', 'Lead Time, Ciclo, Faturamento Minimo', 'Novo v5.0'],
        ['Pedido Manual', 'Entrada manual de pedidos', 'Ativo'],
        ['Central de Parametros', 'Consulta e edicao de parametros', 'Novo v5.0']
    ]
    add_table_with_header(doc, modulos_headers, modulos_data)

    doc.add_page_break()

    # ========================================
    # 3. ARQUITETURA DO SISTEMA
    # ========================================
    doc.add_heading('3. Arquitetura do Sistema', level=1)

    arq_text = """O sistema segue uma arquitetura modular com tres camadas principais:

1. BANCO DE DADOS (PostgreSQL)
   - Historico de vendas
   - Cadastros de produtos e fornecedores
   - Estoques atuais
   - Parametros de configuracao

2. CAMADA DE PREVISAO (Bottom-Up V2)
   - 6 metodos estatisticos
   - Deteccao de sazonalidade
   - Limitadores de tendencia
   - Tratamento de outliers

3. CAMADA DE PEDIDO
   - Cobertura baseada em ABC
   - Estoque de seguranca
   - Arredondamento multiplo de caixa
   - Agregacao por fornecedor/loja"""

    doc.add_paragraph(arq_text)

    doc.add_heading('URLs de Acesso', level=2)

    urls_headers = ['Modulo', 'URL', 'Descricao']
    urls_data = [
        ['Menu Principal', '/menu', 'Pagina inicial com acesso a todos os modulos'],
        ['Previsao de Demanda', '/', 'Geracao de previsoes de vendas'],
        ['Pedido ao Fornecedor', '/pedido_fornecedor_integrado', 'Pedido multi-loja com calculo ABC'],
        ['Transferencias', '/transferencias', 'Gestao de transferencias entre lojas'],
        ['Simulador', '/simulador', 'Simulacao de cenarios what-if'],
        ['Eventos', '/eventos_simples', 'Calendario promocional'],
        ['KPIs e Metricas', '/kpis', 'Dashboard de indicadores'],
        ['Pedido Manual', '/pedido_manual', 'Entrada manual de pedidos'],
        ['Central de Parametros', '/central-parametros-fornecedor', 'Parametros de fornecedor']
    ]
    add_table_with_header(doc, urls_headers, urls_data)

    doc.add_page_break()

    # ========================================
    # 4. INSTALACAO E CONFIGURACAO
    # ========================================
    doc.add_heading('4. Instalacao e Configuracao', level=1)

    doc.add_heading('4.1 Pre-requisitos', level=2)

    prereq = [
        "Python 3.8 ou superior",
        "PostgreSQL 15 ou superior",
        "pip (gerenciador de pacotes Python)"
    ]
    for item in prereq:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.2 Passo a Passo', level=2)

    doc.add_paragraph("1. Clone o repositorio:")
    doc.add_paragraph("   git clone https://github.com/conradocosta0602/previsao-demanda.git")

    doc.add_paragraph("2. Crie o ambiente virtual:")
    doc.add_paragraph("   python -m venv venv")
    doc.add_paragraph("   venv\\Scripts\\activate  (Windows)")

    doc.add_paragraph("3. Instale as dependencias:")
    doc.add_paragraph("   pip install -r requirements.txt")

    doc.add_paragraph("4. Configure o banco de dados PostgreSQL")

    doc.add_paragraph("5. Execute o sistema:")
    doc.add_paragraph("   python app.py")

    doc.add_paragraph("6. Acesse no navegador:")
    doc.add_paragraph("   http://localhost:5001/menu")

    doc.add_heading('4.3 Requisitos Tecnicos', level=2)

    req_headers = ['Componente', 'Versao']
    req_data = [
        ['Python', '3.8+'],
        ['Flask', '3.0.0'],
        ['Pandas', '2.1.3'],
        ['NumPy', '1.26.2'],
        ['SciPy', '1.11.4'],
        ['Scikit-learn', '1.3.2'],
        ['Statsmodels', '0.14.0'],
        ['psycopg2-binary', '2.9.9'],
        ['openpyxl', '3.1.2']
    ]
    add_table_with_header(doc, req_headers, req_data)

    doc.add_page_break()

    # ========================================
    # 5. MODULOS DO SISTEMA
    # ========================================
    doc.add_heading('5. Modulos do Sistema', level=1)

    # 5.1 Previsao de Demanda
    doc.add_heading('5.1 Previsao de Demanda', level=2)

    prev_desc = """O modulo de Previsao de Demanda utiliza a abordagem Bottom-Up V2 que:

1. Busca dados do banco PostgreSQL (historico de vendas)
2. Limpa outliers usando IQR e Z-Score
3. Detecta sazonalidade por autocorrelacao
4. Avalia 6 metodos estatisticos
5. Seleciona o melhor baseado em WMAPE historico
6. Aplica limitadores de tendencia
7. Gera previsao para o periodo solicitado"""
    doc.add_paragraph(prev_desc)

    doc.add_heading('Limitadores de Tendencia', level=3)

    limit_headers = ['Tipo', 'Limite', 'Acao']
    limit_data = [
        ['Queda', '> 40%', 'Limitada a -40%'],
        ['Alta', '> 50%', 'Limitada a +50%']
    ]
    add_table_with_header(doc, limit_headers, limit_data)

    doc.add_paragraph()
    doc.add_heading('Granularidades Suportadas', level=3)

    gran_headers = ['Granularidade', 'Periodos', 'Historico Minimo', 'Uso Recomendado']
    gran_data = [
        ['Mensal', '1-24 meses', '12 meses', 'Planejamento estrategico'],
        ['Semanal', '4-104 semanas', '52 semanas', 'Reabastecimento'],
        ['Diario', '7-365 dias', '365 dias', 'Operacoes day-to-day']
    ]
    add_table_with_header(doc, gran_headers, gran_data)

    doc.add_paragraph()
    doc.add_heading('Saneamento de Rupturas', level=3)

    rup_text = """O sistema detecta automaticamente periodos de ruptura (quando venda=0 E estoque=0) e substitui por valores estimados para evitar distorcao na previsao.

Hierarquia de Estimacao:
1. Interpolacao Sazonal - Usa valor do mesmo periodo do ano anterior
2. Media Adjacentes - Media dos periodos imediatamente antes e depois
3. Mediana do Periodo - Mediana dos periodos com venda > 0"""
    doc.add_paragraph(rup_text)

    doc.add_paragraph()
    doc.add_heading('Metricas de Acuracia', level=3)

    metricas = [
        "WMAPE: Weighted Mean Absolute Percentage Error",
        "BIAS: Tendencia sistematica de erro",
        "YoY: Comparacao com mesmo periodo ano anterior"
    ]
    for m in metricas:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_page_break()

    # 5.2 Pedido ao Fornecedor Integrado
    doc.add_heading('5.2 Pedido ao Fornecedor Integrado', level=2)

    ped_desc = """O modulo de Pedido ao Fornecedor Integrado permite gerar pedidos de compra otimizados considerando:

- Previsao de demanda baseada em historico e sazonalidade
- Cobertura de estoque configuravel por curva ABC
- Parametros de fornecedor (Lead Time, Ciclo de Pedido, Faturamento Minimo)
- Situacao de compra dos itens (bloqueios, restricoes)
- Transferencias entre lojas do mesmo grupo regional"""
    doc.add_paragraph(ped_desc)

    doc.add_heading('Fluxos de Pedido', level=3)

    doc.add_paragraph("1. Pedido para Lojas (Direto)").bold = True
    doc.add_paragraph("Gera pedidos individuais para cada loja selecionada, considerando o estoque e demanda especificos de cada unidade.")

    doc.add_paragraph("2. Pedido para CD (Centralizado)").bold = True
    doc.add_paragraph("Gera um pedido consolidado para o Centro de Distribuicao, agregando a demanda de todas as lojas do grupo regional.")

    doc.add_heading('Formula de Cobertura', level=3)

    doc.add_paragraph("Cobertura (dias) = Lead Time + Ciclo Pedido (7d) + Seguranca ABC")

    cob_headers = ['Curva ABC', 'Seguranca (dias)', 'Exemplo (LT=15d)']
    cob_data = [
        ['A', '+2 dias', '15 + 7 + 2 = 24 dias'],
        ['B', '+4 dias', '15 + 7 + 4 = 26 dias'],
        ['C', '+6 dias', '15 + 7 + 6 = 28 dias']
    ]
    add_table_with_header(doc, cob_headers, cob_data)

    doc.add_paragraph()
    doc.add_heading('Estoque de Seguranca', level=3)

    doc.add_paragraph("ES = Z x Sigma x Raiz(LT)")
    doc.add_paragraph("Onde: Z = Fator Z do nivel de servico, Sigma = Desvio padrao diario, LT = Lead time em dias")

    es_headers = ['Curva ABC', 'Nivel Servico', 'Fator Z']
    es_data = [
        ['A', '98%', '2.05'],
        ['B', '95%', '1.65'],
        ['C', '90%', '1.28']
    ]
    add_table_with_header(doc, es_headers, es_data)

    doc.add_paragraph()
    doc.add_heading('Indicadores Visuais', level=3)

    badge_headers = ['Badge', 'Cor', 'Significado']
    badge_data = [
        ['Pedir', 'Azul', 'Item precisa ser pedido'],
        ['Bloqueado', 'Amarelo', 'Item com restricao de compra'],
        ['OK', 'Verde', 'Estoque suficiente'],
        ['Ruptura', 'Vermelho', 'Ruptura iminente'],
        ['Enviar X un', 'Laranja', 'Loja vai enviar para outra'],
        ['Receber X un', 'Verde', 'Loja vai receber de outra']
    ]
    add_table_with_header(doc, badge_headers, badge_data)

    doc.add_page_break()

    # 5.3 Transferencias entre Lojas
    doc.add_heading('5.3 Transferencias entre Lojas', level=2)

    transf_desc = """O modulo de Transferencias entre Lojas permite:

- Identificar automaticamente desbalanceamentos de estoque
- Sugerir transferencias entre lojas do mesmo grupo
- Priorizar por urgencia (ruptura, alta, media, baixa)
- Exportar lista de transferencias para Excel
- Integrar com o fluxo de Pedido ao Fornecedor (CD)"""
    doc.add_paragraph(transf_desc)

    doc.add_heading('Niveis de Urgencia', level=3)

    urg_headers = ['Urgencia', 'Cobertura Destino', 'Cor', 'Acao']
    urg_data = [
        ['CRITICA', '0 dias (ruptura)', 'Vermelho', 'Prioridade maxima'],
        ['ALTA', '1-3 dias', 'Laranja', 'Atender em ate 24h'],
        ['MEDIA', '4-7 dias', 'Amarelo', 'Atender na proxima rota'],
        ['BAIXA', '> 7 dias', 'Verde', 'Atender quando conveniente']
    ]
    add_table_with_header(doc, urg_headers, urg_data)

    doc.add_paragraph()
    doc.add_heading('Regras de Transferencia', level=3)

    regras = [
        "Loja doadora mantem minimo de 10 dias de cobertura apos doacao",
        "Margem de excesso: 7 dias acima da cobertura alvo",
        "Transferencia nao pode deixar doadora em situacao de ruptura",
        "Mesma loja NAO pode enviar e receber o mesmo produto",
        "Apenas lojas do mesmo grupo podem transferir"
    ]
    for r in regras:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # 5.4 a 5.8 Outros modulos
    doc.add_heading('5.4 Simulador de Cenarios', level=2)
    doc.add_paragraph("O Simulador permite criar cenarios what-if para analisar o impacto de diferentes variaveis na demanda, como promocoes, eventos sazonais e alteracoes de preco.")

    doc.add_heading('5.5 Gerenciamento de Eventos', level=2)
    doc.add_paragraph("O modulo de Eventos permite cadastrar promocoes e eventos que impactam a demanda, permitindo que o sistema considere esses fatores nas previsoes.")

    doc.add_heading('5.6 KPIs e Metricas', level=2)
    doc.add_paragraph("Dashboard com indicadores de performance do sistema, incluindo acuracia das previsoes, cobertura de estoque e nivel de servico.")

    doc.add_heading('5.7 Pedido Manual', level=2)
    doc.add_paragraph("Permite criar pedidos de compra manualmente com controle total sobre quantidades, util para itens especiais ou situacoes excepcionais.")

    doc.add_heading('5.8 Central de Parametros', level=2)
    doc.add_paragraph("Permite consultar e editar parametros de fornecedor como Lead Time, Ciclo de Pedido e Faturamento Minimo, alem de visualizar produtos por fornecedor.")

    doc.add_page_break()

    # ========================================
    # 6. POLITICA DE ESTOQUE ABC
    # ========================================
    doc.add_heading('6. Politica de Estoque ABC', level=1)

    abc_text = """A classificacao ABC e definida no cadastro de produtos e reflete a importancia de cada item no faturamento."""
    doc.add_paragraph(abc_text)

    abc_headers = ['Curva', 'Caracteristica', '% SKUs (tipico)', '% Faturamento (tipico)']
    abc_data = [
        ['A', 'Alto giro', '20%', '80%'],
        ['B', 'Medio giro', '30%', '15%'],
        ['C', 'Baixo giro', '50%', '5%']
    ]
    add_table_with_header(doc, abc_headers, abc_data)

    doc.add_paragraph()
    doc.add_heading('Impacto da Curva ABC', level=2)

    impacto = [
        "Nivel de Servico: A=98%, B=95%, C=90%",
        "Dias de Seguranca: A=+2d, B=+4d, C=+6d",
        "Prioridade de Atencao: Itens A recebem mais foco"
    ]
    for i in impacto:
        doc.add_paragraph(i, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # 7. APIS DISPONIVEIS
    # ========================================
    doc.add_heading('7. APIs Disponiveis', level=1)

    api_headers = ['Endpoint', 'Metodo', 'Descricao']
    api_data = [
        ['/api/gerar_previsao_banco', 'POST', 'Gera previsao de demanda'],
        ['/api/pedido_fornecedor_integrado', 'POST', 'Calcula pedido ao fornecedor'],
        ['/api/transferencias/oportunidades', 'GET', 'Lista oportunidades de transferencia'],
        ['/api/transferencias/grupos', 'GET', 'Lista grupos de transferencia'],
        ['/api/transferencias/exportar', 'GET', 'Exporta transferencias para Excel'],
        ['/api/demanda_validada/salvar', 'POST', 'Salva demanda validada'],
        ['/api/demanda_validada/listar', 'GET', 'Lista demandas validadas'],
        ['/api/fornecedores', 'GET', 'Lista fornecedores cadastrados'],
        ['/api/lojas', 'GET', 'Lista lojas cadastradas']
    ]
    add_table_with_header(doc, api_headers, api_data)

    doc.add_page_break()

    # ========================================
    # 8. ESTRUTURA DO BANCO DE DADOS
    # ========================================
    doc.add_heading('8. Estrutura do Banco de Dados', level=1)

    doc.add_heading('Principais Tabelas', level=2)

    tabelas = [
        ("historico_vendas_diario", "Historico de vendas por dia, produto e empresa"),
        ("cadastro_produtos", "Cadastro de produtos com curva ABC"),
        ("cadastro_fornecedores", "Cadastro de fornecedores com lead time"),
        ("estoque_atual", "Estoque disponivel por produto e empresa"),
        ("parametros_gondola", "Parametros de multiplo e lote minimo"),
        ("parametros_fornecedor", "Parametros especificos por fornecedor"),
        ("grupos_transferencia", "Grupos regionais de lojas"),
        ("lojas_grupo_transferencia", "Lojas de cada grupo"),
        ("oportunidades_transferencia", "Historico de transferencias"),
        ("situacao_compra_regras", "Regras de bloqueio de compra"),
        ("situacao_compra_itens", "Itens bloqueados")
    ]

    for nome, desc in tabelas:
        p = doc.add_paragraph()
        run = p.add_run(nome + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========================================
    # 9. FAQ
    # ========================================
    doc.add_heading('9. FAQ - Perguntas Frequentes', level=1)

    faqs = [
        ("Qual a diferenca entre AUTO e ML na selecao de metodo?",
         "AUTO usa regras baseadas em caracteristicas da demanda (CV, tendencia, sazonalidade). ML usa Random Forest com 15+ features. ML tende a ser mais preciso, AUTO e mais rapido."),

        ("Quantos meses de historico preciso?",
         "Mensal: Minimo 12 meses, ideal 24+ meses. Semanal: Minimo 52 semanas, ideal 104 semanas. Diario: Minimo 365 dias, ideal 730 dias."),

        ("O que significam os limitadores de tendencia?",
         "Quedas >40% sao limitadas a -40% (protecao contra dados anomalos). Altas >50% sao limitadas a +50% (protecao contra otimismo excessivo)."),

        ("Como a ferramenta trata rupturas no historico?",
         "O sistema detecta automaticamente rupturas (venda=0 E estoque=0) e substitui por valores estimados usando interpolacao sazonal."),

        ("Por que a cobertura e calculada com ABC e nao com CV?",
         "A previsao V2 ja captura a variabilidade da demanda atraves dos 6 metodos estatisticos, limitadores e deteccao de outliers. Adicionar CV seria redundante."),

        ("Posso usar uma cobertura fixa?",
         "Sim. O filtro 'Cobertura (dias)' permite escolher entre automatica (ABC) ou valores fixos de 15, 21, 30 ou 45 dias."),

        ("O que acontece se um produto nao tem curva ABC cadastrada?",
         "O sistema assume curva 'B' como padrao (nivel de servico 95%, +4 dias de seguranca).")
    ]

    for pergunta, resposta in faqs:
        p = doc.add_paragraph()
        run = p.add_run("P: " + pergunta)
        run.bold = True
        doc.add_paragraph("R: " + resposta)
        doc.add_paragraph()

    doc.add_page_break()

    # ========================================
    # 10. CHANGELOG
    # ========================================
    doc.add_heading('10. Changelog', level=1)

    doc.add_heading('v5.0 (Janeiro 2026) - ATUAL', level=2)

    doc.add_paragraph("Novas Funcionalidades:").bold = True
    v5_new = [
        "Pedido Multi-Loja: Layout hierarquico Fornecedor > Loja > Itens com multi-selecao",
        "Transferencias entre Lojas: Identificacao automatica de oportunidades de balanceamento",
        "Parametros de Fornecedor: Importacao de Lead Time, Ciclo de Pedido e Faturamento Minimo",
        "Graceful Degradation: Sistema funciona mesmo sem tabelas novas",
        "Central de Parametros: Consulta e edicao de parametros de fornecedor",
        "Coluna Situacao de Compra na previsao de demanda"
    ]
    for item in v5_new:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("Melhorias de Interface:").bold = True
    v5_melhoria = [
        "Multi-selecao em filtros (Lojas, Fornecedores, Categorias)",
        "Badges visuais para transferencias (Enviar/Receber)",
        "Ordenacao do Excel por Codigo Filial e CNPJ Fornecedor",
        "Botao 'Limpar Historico' na tela de transferencias",
        "Ordenacao por criticidade de alertas na tabela de previsao detalhada",
        "Congelamento das colunas Codigo e Descricao no relatorio detalhado"
    ]
    for item in v5_melhoria:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("Remocoes:").bold = True
    doc.add_paragraph("Upload de arquivo Excel na tela de Previsao de Demanda (agora somente via banco de dados)", style='List Bullet')

    doc.add_heading('v4.0 (Janeiro 2026)', level=2)
    v4 = [
        "Modulo de Pedido ao Fornecedor Integrado com previsao V2",
        "Politica de estoque baseada exclusivamente em curva ABC",
        "Formula de cobertura simplificada",
        "Interface unificada no menu principal"
    ]
    for item in v4:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('v3.0 (Dezembro 2025)', level=2)
    v3 = [
        "Previsao de demanda com 6 metodos",
        "Selecao automatica (AUTO) e Machine Learning (ML)",
        "Reabastecimento com 3 fluxos",
        "Calendario promocional",
        "Simulador de cenarios"
    ]
    for item in v3:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # 11. SUPORTE E CONTATO
    # ========================================
    doc.add_heading('11. Suporte e Contato', level=1)

    doc.add_heading('Autores', level=2)
    autores = [
        "Valter Lino (@valterlino01) - Desenvolvedor Principal",
        "Conrado Costa (@conradocosta0602) - Co-desenvolvedor",
        "Consultoria Tecnica: Claude (Anthropic)"
    ]
    for a in autores:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_heading('Suporte', level=2)
    doc.add_paragraph("Para duvidas ou problemas:")
    doc.add_paragraph("- Consulte a documentacao completa no GitHub", style='List Bullet')
    doc.add_paragraph("- Abra uma issue em: https://github.com/conradocosta0602/previsao-demanda/issues", style='List Bullet')

    doc.add_heading('Features Planejadas', level=2)
    features = [
        "API REST para integracao com ERP",
        "Dashboard Power BI",
        "Pedido CD -> Lojas integrado",
        "Transferencias automaticas",
        "Exportacao de graficos",
        "Modo batch para grandes volumes"
    ]
    for f in features:
        doc.add_paragraph("[ ] " + f, style='List Bullet')

    # Rodape final
    doc.add_paragraph()
    doc.add_paragraph()
    final = doc.add_paragraph()
    final_run = final.add_run("Sistema de Demanda e Reabastecimento v5.0")
    final_run.bold = True
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status = doc.add_paragraph()
    status.add_run("Status: Em Producao | Ultima Atualizacao: Janeiro 2026")
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salvar documento
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               'Documentacao_Sistema_Demanda_Reabastecimento_v5.docx')
    doc.save(output_path)
    print(f"Documento salvo em: {output_path}")
    return output_path

if __name__ == "__main__":
    criar_documento()
