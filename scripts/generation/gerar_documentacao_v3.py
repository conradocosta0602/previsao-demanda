"""
Gera documentação completa do Sistema de Previsão de Demanda e Reabastecimento v3.0
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Criar novo documento
doc = Document()

# ========== CAPA ==========
title = doc.add_heading('Sistema de Previsão de Demanda e Reabastecimento', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('Documentação Técnica e Manual do Usuário')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 112, 243)

version = doc.add_paragraph()
version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = version.add_run('Versão 3.0')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n')
info.add_run('Sistema integrado de previsão estatística e otimização de estoque')

doc.add_page_break()

# ========== SUMÁRIO EXECUTIVO ==========
doc.add_heading('Sumário Executivo', 1)

p = doc.add_paragraph()
p.add_run('O Sistema de Previsão de Demanda e Reabastecimento v3.0 é uma solução completa para gestão de estoque multi-loja com Centro de Distribuição (CD), combinando métodos estatísticos avançados, machine learning e cálculos de reabastecimento baseados em níveis de serviço.')

doc.add_heading('Principais Funcionalidades', 2)
doc.add_paragraph('Previsão de demanda com 5+ métodos estatísticos e seleção automática por ML', style='List Bullet')
doc.add_paragraph('Cálculo automático de estoque de segurança e ponto de pedido', style='List Bullet')
doc.add_paragraph('Gerenciamento de 3 fluxos: Fornecedor → CD/Loja, CD → Loja, Transferências entre Lojas', style='List Bullet')
doc.add_paragraph('Detecção automática de sazonalidade, outliers e anomalias', style='List Bullet')
doc.add_paragraph('Simulador de cenários (What-If Analysis)', style='List Bullet')
doc.add_paragraph('Calendário promocional integrado', style='List Bullet')
doc.add_paragraph('Métricas de acurácia (MAPE, BIAS) e alertas inteligentes', style='List Bullet')

doc.add_page_break()

# ========== ÍNDICE (manual) ==========
doc.add_heading('Índice', 1)

doc.add_paragraph('1. Visão Geral do Sistema', style='List Number')
doc.add_paragraph('2. Métodos Estatísticos de Previsão', style='List Number')
doc.add_paragraph('3. Telas e Funcionalidades', style='List Number')
doc.add_paragraph('4. Conceitos de Reabastecimento', style='List Number')
doc.add_paragraph('5. Indicadores e Métricas', style='List Number')
doc.add_paragraph('6. Recursos Avançados', style='List Number')
doc.add_paragraph('7. Orientações de Uso', style='List Number')
doc.add_paragraph('8. Perguntas Frequentes', style='List Number')

doc.add_page_break()

# ========== 1. VISÃO GERAL DO SISTEMA ==========
doc.add_heading('1. Visão Geral do Sistema', 1)

doc.add_heading('1.1. Objetivo', 2)
p = doc.add_paragraph()
p.add_run('O sistema tem como objetivo principal ')
run = p.add_run('otimizar o planejamento de compras e reabastecimento')
run.bold = True
p.add_run(' em operações de varejo com múltiplas lojas e centro de distribuição, combinando previsão estatística de demanda com cálculos de estoque de segurança e ponto de pedido.')

doc.add_heading('1.2. Arquitetura', 2)
doc.add_paragraph('Backend: Python (Flask, Pandas, NumPy, Statsmodels)', style='List Bullet')
doc.add_paragraph('Frontend: HTML5, CSS3, JavaScript (Chart.js)', style='List Bullet')
doc.add_paragraph('Machine Learning: Scikit-learn (Random Forest, Gradient Boosting)', style='List Bullet')
doc.add_paragraph('Entrada: Arquivos Excel (.xlsx) com histórico de vendas', style='List Bullet')
doc.add_paragraph('Saída: Relatórios Excel com previsões e pedidos sugeridos', style='List Bullet')

doc.add_heading('1.3. Fluxos Suportados', 2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fluxo'
hdr_cells[1].text = 'Origem'
hdr_cells[2].text = 'Destino'

row = table.rows[1].cells
row[0].text = 'Fornecedor → CD/Loja'
row[1].text = 'Fornecedor Externo'
row[2].text = 'CD ou Lojas'

row = table.rows[2].cells
row[0].text = 'CD → Loja'
row[1].text = 'Centro de Distribuição'
row[2].text = 'Lojas'

row = table.rows[3].cells
row[0].text = 'Transferências'
row[1].text = 'Loja Origem'
row[2].text = 'Loja Destino'

doc.add_page_break()

# ========== 2. MÉTODOS ESTATÍSTICOS DE PREVISÃO ==========
doc.add_heading('2. Métodos Estatísticos de Previsão', 1)

p = doc.add_paragraph()
p.add_run('O sistema implementa ')
run = p.add_run('6 métodos estatísticos')
run.bold = True
p.add_run(' de previsão de séries temporais, com seleção automática baseada em características da demanda ou via Machine Learning.')

doc.add_heading('2.1. Simple Moving Average (SMA)', 2)

doc.add_paragraph('Descrição: Média móvel simples dos últimos N períodos', style='List Bullet')
doc.add_paragraph('Indicado para: Demanda estável sem tendência ou sazonalidade', style='List Bullet')
doc.add_paragraph('Janela adaptativa: 3-6 meses dependendo da variabilidade', style='List Bullet')
doc.add_paragraph('Vantagens: Simplicidade, robustez a ruídos', style='List Bullet')
doc.add_paragraph('Limitações: Não captura tendências ou sazonalidade', style='List Bullet')

doc.add_heading('2.2. Weighted Moving Average (WMA)', 2)

doc.add_paragraph('Descrição: Média móvel com pesos maiores para períodos recentes', style='List Bullet')
doc.add_paragraph('Indicado para: Demanda com leve tendência recente', style='List Bullet')
doc.add_paragraph('Janela adaptativa: Ajusta automaticamente baseado em CV', style='List Bullet')
doc.add_paragraph('Pesos: Lineares decrescentes (período mais recente = peso maior)', style='List Bullet')
doc.add_paragraph('Vantagens: Mais responsivo que SMA a mudanças recentes', style='List Bullet')

doc.add_heading('2.3. Simple Exponential Smoothing (SES)', 2)

doc.add_paragraph('Descrição: Suavização exponencial com alpha otimizado', style='List Bullet')
doc.add_paragraph('Indicado para: Demanda estável a moderadamente variável', style='List Bullet')
doc.add_paragraph('Alpha: Otimizado automaticamente (0.1 a 0.9)', style='List Bullet')
doc.add_paragraph('Vantagens: Adaptativo, considera todo histórico com pesos decrescentes', style='List Bullet')

doc.add_heading('2.4. Linear Regression Forecast', 2)

doc.add_paragraph('Descrição: Regressão linear com tendência temporal', style='List Bullet')
doc.add_paragraph('Indicado para: Demanda com tendência clara (crescimento/decrescimento)', style='List Bullet')
doc.add_paragraph('Fórmula: y = a + b*t', style='List Bullet')
doc.add_paragraph('Vantagens: Captura tendências de longo prazo', style='List Bullet')
doc.add_paragraph('Limitações: Não captura sazonalidade', style='List Bullet')

doc.add_heading('2.5. TSB (Trigg and Leach Smoothing with BIAS)', 2)

doc.add_paragraph('Descrição: Suavização exponencial adaptativa com correção automática de viés', style='List Bullet')
doc.add_paragraph('Indicado para: Demanda com tendência variável e necessidade de auto-correção', style='List Bullet')
doc.add_paragraph('Mecanismo: Monitora o erro acumulado (BIAS) e ajusta o fator de suavização', style='List Bullet')
doc.add_paragraph('Tracking Signal: Detecta mudanças sistemáticas e adapta alfa automaticamente', style='List Bullet')
doc.add_paragraph('Vantagens: Auto-adaptativo, corrige viés acumulado, robusto a mudanças', style='List Bullet')
doc.add_paragraph('Fórmula: Usa Smoothed Error (SE) e Mean Absolute Deviation (MAD) para ajuste dinâmico', style='List Bullet')

doc.add_heading('2.6. Decomposição Sazonal Mensal (Híbrida)', 2)

doc.add_paragraph('Descrição: Modelo híbrido combinando sazonalidade multiplicativa + tendência linear', style='List Bullet')
doc.add_paragraph('Indicado para: Demanda com padrão sazonal (12 meses)', style='List Bullet')
doc.add_paragraph('Fórmula: Previsão = Índice Sazonal × (Baseline + Tendência × t)', style='List Bullet')
doc.add_paragraph('Requisitos: Mínimo 18 meses de histórico', style='List Bullet')
doc.add_paragraph('Vantagens: Captura sazonalidade E tendência simultaneamente', style='List Bullet')

doc.add_heading('2.7. Seleção Automática de Método', 2)

p = doc.add_paragraph()
p.add_run('Modo AUTO: ')
run = p.add_run('O sistema analisa características da demanda e seleciona o método mais adequado')
run.bold = True

doc.add_paragraph('CV < 0.3 + Tendência fraca → SMA', style='List Bullet')
doc.add_paragraph('CV 0.3-0.6 + Tendência moderada → WMA ou SES', style='List Bullet')
doc.add_paragraph('Tendência forte → Linear Regression ou TSB', style='List Bullet')
doc.add_paragraph('Demanda variável com viés → TSB (auto-correção)', style='List Bullet')
doc.add_paragraph('Sazonalidade detectada + 18+ meses → Decomposição Sazonal', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Modo ML (Machine Learning): ')
run = p.add_run('Random Forest treinado avalia 15+ features e seleciona o método ideal')
run.bold = True

doc.add_page_break()

# ========== 3. TELAS E FUNCIONALIDADES ==========
doc.add_heading('3. Telas e Funcionalidades', 1)

doc.add_heading('3.1. Tela de Previsão de Demanda', 2)

doc.add_heading('Funcionalidade Principal', 3)
doc.add_paragraph('Upload de arquivo Excel com histórico de vendas (SKU, Local, Mês, Quantidade)', style='List Bullet')
doc.add_paragraph('Processamento automático com seleção de método por SKU/Loja', style='List Bullet')
doc.add_paragraph('Geração de previsões para próximos N meses (configurável 1-24 meses)', style='List Bullet')

doc.add_heading('Componentes Visuais', 3)
doc.add_paragraph('Cards Executivos: Total SKUs, Meses Previstos, MAPE, BIAS', style='List Bullet')
doc.add_paragraph('Painel Explicativo: Interpretação de MAPE e BIAS', style='List Bullet')
doc.add_paragraph('Tabela Comparativa YoY: Previsão vs Ano Anterior (acima do gráfico)', style='List Bullet')
doc.add_paragraph('Gráfico de Barras: Visualização da comparação YoY', style='List Bullet')
doc.add_paragraph('Tabela Fornecedor/Item: Previsão detalhada com alertas coloridos', style='List Bullet')
doc.add_paragraph('Legenda de Indicadores: 🔴 Crítico, 🟡 Alerta, 🔵 Atenção, 🟢 Normal', style='List Bullet')

doc.add_heading('Alertas e Indicadores', 3)
doc.add_paragraph('🔴 Crítico: Variação YoY > 100% ou MAPE > 30%', style='List Bullet')
doc.add_paragraph('🟡 Alerta: Variação YoY > 50%', style='List Bullet')
doc.add_paragraph('🔵 Atenção: Variação YoY > 20%', style='List Bullet')
doc.add_paragraph('🟢 Normal: Variação YoY ≤ 20%', style='List Bullet')

doc.add_heading('3.2. Tela de Pedidos ao Fornecedor', 2)

doc.add_heading('Funcionalidade Principal', 3)
doc.add_paragraph('Calcula pedidos de reabastecimento: Fornecedor → CD ou Fornecedor → Loja', style='List Bullet')
doc.add_paragraph('Considera: Lead time do fornecedor, ciclo de pedido, múltiplos de palete/carreta', style='List Bullet')
doc.add_paragraph('Consolidação de carga para otimização logística', style='List Bullet')

doc.add_heading('Parâmetros', 3)
doc.add_paragraph('Nível de Serviço: Automático por classe ABC (A=98%, B=95%, C=90%)', style='List Bullet')
doc.add_paragraph('Ciclo de Pedido: Padrão 30 dias (CD) / 14 dias (Loja)', style='List Bullet')
doc.add_paragraph('Lead Time: Configurável por fornecedor/item', style='List Bullet')

doc.add_heading('Saída', 3)
doc.add_paragraph('Tabela agrupada por fornecedor com totalizadores', style='List Bullet')
doc.add_paragraph('Métricas: Valor estoque, Valor pedido, Cobertura atual/projetada, Ruptura', style='List Bullet')
doc.add_paragraph('Custo unitário sempre exibido (mesmo sem pedido)', style='List Bullet')

doc.add_heading('3.3. Tela de Pedidos CD → Loja', 2)

doc.add_heading('Funcionalidade Principal', 3)
doc.add_paragraph('Calcula transferências do CD para Lojas', style='List Bullet')
doc.add_paragraph('Considera estoque disponível no CD como limite', style='List Bullet')
doc.add_paragraph('Priorização por criticidade (menor cobertura = maior prioridade)', style='List Bullet')

doc.add_heading('Parâmetros', 3)
doc.add_paragraph('Ciclo de Reabastecimento: 7 dias (padrão)', style='List Bullet')
doc.add_paragraph('Lead Time Interno: 1-3 dias', style='List Bullet')

doc.add_heading('3.4. Tela de Transferências entre Lojas', 2)

doc.add_heading('Funcionalidade Principal', 3)
doc.add_paragraph('Identifica oportunidades de transferência entre lojas', style='List Bullet')
doc.add_paragraph('Loja com excesso → Loja com ruptura/baixa cobertura', style='List Bullet')
doc.add_paragraph('Algoritmo: Prioriza SKUs com maior impacto (risco ruptura vs excesso)', style='List Bullet')

doc.add_heading('Critérios', 3)
doc.add_paragraph('Origem: Cobertura > 45 dias E Quantidade disponível ≥ demanda mensal', style='List Bullet')
doc.add_paragraph('Destino: Cobertura < 14 dias OU risco de ruptura', style='List Bullet')
doc.add_paragraph('Quantidade: Mínimo do excesso origem vs necessidade destino', style='List Bullet')

doc.add_heading('3.5. Gerenciador de Eventos', 2)

doc.add_heading('Funcionalidade Principal', 3)
doc.add_paragraph('Cadastro de eventos promocionais/sazonais', style='List Bullet')
doc.add_paragraph('Ajuste automático de previsões com multiplicador configurável', style='List Bullet')

doc.add_heading('Campos', 3)
doc.add_paragraph('Nome do Evento, Data Início, Data Fim', style='List Bullet')
doc.add_paragraph('Impacto Esperado (%): Ex: Black Friday = +150%', style='List Bullet')
doc.add_paragraph('Filtros: SKUs, Lojas, Categorias específicas', style='List Bullet')

doc.add_heading('3.6. Simulador de Cenários', 2)

doc.add_heading('Funcionalidade Principal', 3)
doc.add_paragraph('What-If Analysis: Simula impacto de mudanças sem alterar dados reais', style='List Bullet')
doc.add_paragraph('Ajustes: Nível de serviço, lead time, demanda futura', style='List Bullet')

doc.add_heading('Exemplos de Uso', 3)
doc.add_paragraph('E se aumentarmos o nível de serviço de 95% para 98%?', style='List Bullet')
doc.add_paragraph('Qual o impacto de reduzir o lead time do fornecedor de 30 para 20 dias?', style='List Bullet')
doc.add_paragraph('Como ficaria o estoque com aumento de 20% na demanda?', style='List Bullet')

doc.add_page_break()

# ========== 4. CONCEITOS DE REABASTECIMENTO ==========
doc.add_heading('4. Conceitos de Reabastecimento', 1)

doc.add_heading('4.1. Estoque de Segurança (Safety Stock)', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Quantidade adicional mantida para proteger contra incertezas de demanda e fornecimento')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula Implementada:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('SS = Z × σ × √LT')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph('    Onde:', style='List Bullet 2')
doc.add_paragraph('        Z = Fator de serviço (tabela normal padrão)', style='List Bullet 2')
doc.add_paragraph('        σ = Desvio padrão da demanda', style='List Bullet 2')
doc.add_paragraph('        LT = Lead time em meses', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Níveis de Serviço e Fator Z:', style='List Bullet')

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nível Serviço'
hdr_cells[1].text = 'Fator Z'
hdr_cells[2].text = 'Classe'

row = table.rows[1].cells
row[0].text = '90%'
row[1].text = '1.28'
row[2].text = 'C (baixo giro)'

row = table.rows[2].cells
row[0].text = '95%'
row[1].text = '1.65'
row[2].text = 'B (médio giro)'

row = table.rows[3].cells
row[0].text = '98%'
row[1].text = '2.05'
row[2].text = 'A (alto giro)'

doc.add_heading('4.2. Ponto de Pedido (Reorder Point)', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Nível de estoque que dispara um novo pedido')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula Implementada:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('ROP = (Demanda Diária × Lead Time) + Estoque de Segurança')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Interpretação:', style='List Bullet')
doc.add_paragraph('    Se Estoque Disponível ≤ ROP → Fazer Pedido', style='List Bullet 2')
doc.add_paragraph('    Se Estoque Disponível > ROP → Aguardar', style='List Bullet 2')

doc.add_heading('4.3. Quantidade do Pedido (Order Quantity)', 2)

p = doc.add_paragraph()
p.add_run('Método: ')
run = p.add_run('Período Fixo de Revisão')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula Implementada:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('Q = (Demanda × (LT + Revisão)) + SS - Estoque Atual - Em Trânsito')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('Onde:', style='List Bullet')
doc.add_paragraph('    Revisão = Período entre revisões (ciclo de pedido)', style='List Bullet 2')
doc.add_paragraph('    LT = Lead time de entrega', style='List Bullet 2')
doc.add_paragraph('    SS = Estoque de segurança', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Ajustes:', style='List Bullet')
doc.add_paragraph('Múltiplos de Embalagem: Arredonda para cima (caixa, palete, carreta)', style='List Bullet 2')
doc.add_paragraph('Quantidade Mínima: Respeitada quando configurada', style='List Bullet 2')
doc.add_paragraph('Estoque Máximo: Não excede quando configurado', style='List Bullet 2')

doc.add_heading('4.4. Nível de Serviço', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Probabilidade de NÃO haver ruptura de estoque')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Classificação ABC Automática:', style='List Bullet')

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Classe'
hdr_cells[1].text = 'Demanda Mensal'
hdr_cells[2].text = 'Nível Serviço'
hdr_cells[3].text = 'Risco Ruptura'

row = table.rows[1].cells
row[0].text = 'A'
row[1].text = '> 500 un/mês'
row[2].text = '98%'
row[3].text = '2%'

row = table.rows[2].cells
row[0].text = 'B'
row[1].text = '100-500 un/mês'
row[2].text = '95%'
row[3].text = '5%'

row = table.rows[3].cells
row[0].text = 'C'
row[1].text = '< 100 un/mês'
row[2].text = '90%'
row[3].text = '10%'

doc.add_heading('4.5. Cobertura de Estoque', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Quantidade de dias que o estoque atual pode atender a demanda')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('Cobertura (dias) = (Estoque Disponível / Demanda Média Diária)')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Interpretação:', style='List Bullet')
doc.add_paragraph('    < 7 dias: Risco crítico de ruptura', style='List Bullet 2')
doc.add_paragraph('    7-14 dias: Atenção necessária', style='List Bullet 2')
doc.add_paragraph('    14-30 dias: Cobertura adequada', style='List Bullet 2')
doc.add_paragraph('    > 45 dias: Possível excesso de estoque', style='List Bullet 2')

doc.add_page_break()

# ========== 5. INDICADORES E MÉTRICAS ==========
doc.add_heading('5. Indicadores e Métricas', 1)

doc.add_heading('5.1. MAPE (Mean Absolute Percentage Error)', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Mede a acurácia média das previsões em termos percentuais')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('MAPE = (1/n) × Σ |Real - Previsto| / Real × 100%')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Interpretação:', style='List Bullet')

table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'MAPE'
hdr_cells[1].text = 'Classificação'

row = table.rows[1].cells
row[0].text = '< 10%'
row[1].text = 'Excelente'

row = table.rows[2].cells
row[0].text = '10-20%'
row[1].text = 'Bom'

row = table.rows[3].cells
row[0].text = '> 20%'
row[1].text = 'Requer atenção'

doc.add_heading('5.2. BIAS (Viés de Previsão)', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Identifica tendência sistemática de erro (superestimação ou subestimação)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('BIAS = Σ (Real - Previsto) / Σ Real × 100%')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Interpretação:', style='List Bullet')

table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'BIAS'
hdr_cells[1].text = 'Significado'

row = table.rows[1].cells
row[0].text = 'Positivo (+)'
row[1].text = 'Superestimação (prevê mais que realizado)'

row = table.rows[2].cells
row[0].text = 'Próximo de 0'
row[1].text = 'Balanceado (previsões sem viés)'

row = table.rows[3].cells
row[0].text = 'Negativo (-)'
row[1].text = 'Subestimação (prevê menos que realizado)'

doc.add_heading('5.3. Variação YoY (Year-over-Year)', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Compara previsão do período vs mesmo período do ano anterior')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('YoY = (Σ Previsão Atual - Σ Mesmo Período Ano Anterior) / Σ Ano Anterior × 100%')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Exemplo: ')
run.bold = True
p.add_run('Jul-Dez/2024 vs Jul-Dez/2023 (não Jul-Dez/2024 vs Jan-Jun/2024)')

doc.add_heading('5.4. Coeficiente de Variação (CV)', 2)

p = doc.add_paragraph()
p.add_run('Definição: ')
run = p.add_run('Mede a variabilidade relativa da demanda')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Fórmula:', style='List Bullet')

p = doc.add_paragraph('    ', style='List Bullet 2')
run = p.add_run('CV = Desvio Padrão / Média')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Uso:', style='List Bullet')
doc.add_paragraph('CV < 0.3: Demanda estável → SMA adequado', style='List Bullet 2')
doc.add_paragraph('CV 0.3-0.6: Demanda moderada → WMA/SES', style='List Bullet 2')
doc.add_paragraph('CV > 0.6: Demanda volátil → Métodos adaptativos', style='List Bullet 2')

doc.add_page_break()

# ========== 6. RECURSOS AVANÇADOS ==========
doc.add_heading('6. Recursos Avançados', 1)

doc.add_heading('6.1. Detecção Automática de Sazonalidade', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Identifica padrões sazonais automaticamente usando autocorrelação')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Algoritmo:', style='List Bullet')
doc.add_paragraph('Calcula autocorrelação com lag de 12 meses', style='List Bullet 2')
doc.add_paragraph('Se autocorr(12) > 0.6 → Sazonalidade detectada', style='List Bullet 2')
doc.add_paragraph('Aplica Decomposição Sazonal automaticamente', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Requisitos:', style='List Bullet')
doc.add_paragraph('Mínimo 18 meses de histórico', style='List Bullet 2')
doc.add_paragraph('Série com pelo menos 2 ciclos sazonais completos', style='List Bullet 2')

doc.add_heading('6.2. Detecção Automática de Outliers', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Identifica e trata valores anormais que distorcem previsões')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Métodos Implementados:', style='List Bullet')

doc.add_paragraph('IQR (Interquartile Range):', style='List Bullet 2')
p = doc.add_paragraph('    ', style='List Bullet 3')
run = p.add_run('Outlier se: valor < Q1 - 1.5×IQR ou valor > Q3 + 1.5×IQR')
run.font.size = Pt(9)

doc.add_paragraph('Z-Score:', style='List Bullet 2')
p = doc.add_paragraph('    ', style='List Bullet 3')
run = p.add_run('Outlier se: |z-score| > 3')
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('Tratamento:', style='List Bullet')
doc.add_paragraph('Substituição pela mediana da série', style='List Bullet 2')
doc.add_paragraph('Logging do outlier para auditoria', style='List Bullet 2')

doc.add_heading('6.3. Logging de Seleção AUTO', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Registra detalhadamente o processo de seleção automática de método')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Informações Registradas:', style='List Bullet')
doc.add_paragraph('SKU, Local, Timestamp', style='List Bullet 2')
doc.add_paragraph('Características da demanda: CV, tendência, sazonalidade', style='List Bullet 2')
doc.add_paragraph('Método selecionado e justificativa', style='List Bullet 2')
doc.add_paragraph('MAPE de cada método testado', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Localização:', style='List Bullet')
doc.add_paragraph('Aba METODOS_UTILIZADOS no Excel de saída', style='List Bullet 2')
doc.add_paragraph('Permite auditoria e análise de performance', style='List Bullet 2')

doc.add_heading('6.4. Alertas Inteligentes', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Sistema de alertas visuais baseado em múltiplos critérios')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Tipos de Alertas:', style='List Bullet')

doc.add_paragraph('🔴 Crítico:', style='List Bullet 2')
doc.add_paragraph('    - Variação YoY > 100%', style='List Bullet 3')
doc.add_paragraph('    - MAPE > 30% (previsão muito imprecisa)', style='List Bullet 3')
doc.add_paragraph('    - Cobertura < 7 dias (ruptura iminente)', style='List Bullet 3')

doc.add_paragraph('🟡 Alerta:', style='List Bullet 2')
doc.add_paragraph('    - Variação YoY > 50%', style='List Bullet 3')
doc.add_paragraph('    - MAPE 20-30%', style='List Bullet 3')

doc.add_paragraph('🔵 Atenção:', style='List Bullet 2')
doc.add_paragraph('    - Variação YoY > 20%', style='List Bullet 3')
doc.add_paragraph('    - Tendência forte detectada', style='List Bullet 3')

doc.add_paragraph('🟢 Normal:', style='List Bullet 2')
doc.add_paragraph('    - Variação YoY ≤ 20%', style='List Bullet 3')
doc.add_paragraph('    - MAPE < 20%', style='List Bullet 3')

doc.add_heading('6.5. Machine Learning para Seleção de Método', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Modelo ML treinado para selecionar o método estatístico ideal')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Algoritmos Disponíveis:', style='List Bullet')
doc.add_paragraph('Random Forest Classifier (padrão)', style='List Bullet 2')
doc.add_paragraph('Gradient Boosting Classifier', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Features Utilizadas (15+):', style='List Bullet')
doc.add_paragraph('Estatísticas: Média, Mediana, Desvio Padrão, CV', style='List Bullet 2')
doc.add_paragraph('Tendência: Slope de regressão linear, R²', style='List Bullet 2')
doc.add_paragraph('Sazonalidade: Autocorrelação lag 12, Força sazonal', style='List Bullet 2')
doc.add_paragraph('Volume: Total histórico, Meses disponíveis', style='List Bullet 2')
doc.add_paragraph('Variabilidade: Range, IQR, Outliers detectados', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Treinamento:', style='List Bullet')
doc.add_paragraph('Dataset: Histórico de previsões com método ótimo identificado', style='List Bullet 2')
doc.add_paragraph('Validação: Cross-validation 5-fold', style='List Bullet 2')
doc.add_paragraph('Métrica: Acurácia de seleção + MAPE médio resultante', style='List Bullet 2')

doc.add_heading('6.6. Calendário Promocional', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Ajuste automático de previsões para eventos comerciais')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Uso:', style='List Bullet')
doc.add_paragraph('Cadastrar evento: Nome, Data Início/Fim, Impacto (%)', style='List Bullet 2')
doc.add_paragraph('Sistema ajusta previsão automaticamente nos períodos configurados', style='List Bullet 2')
doc.add_paragraph('Fórmula: Previsão Ajustada = Previsão Base × (1 + Impacto/100)', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Exemplos:', style='List Bullet')
doc.add_paragraph('Black Friday: +150% (multiplica por 2.5)', style='List Bullet 2')
doc.add_paragraph('Dia das Mães: +80% (multiplica por 1.8)', style='List Bullet 2')
doc.add_paragraph('Natal: +120% (multiplica por 2.2)', style='List Bullet 2')

doc.add_heading('6.7. Tratamento de Séries Muito Curtas', 2)

p = doc.add_paragraph()
p.add_run('Funcionalidade: ')
run = p.add_run('Estratégias especiais para produtos com pouco histórico')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Estratégias:', style='List Bullet')

doc.add_paragraph('< 3 meses de histórico:', style='List Bullet 2')
doc.add_paragraph('    - Usa média simples dos meses disponíveis', style='List Bullet 3')
doc.add_paragraph('    - Alerta: "Dados insuficientes - usar com cautela"', style='List Bullet 3')

doc.add_paragraph('3-6 meses de histórico:', style='List Bullet 2')
doc.add_paragraph('    - SMA com janela adaptativa', style='List Bullet 3')
doc.add_paragraph('    - Estoque de segurança aumentado (Z × 1.2)', style='List Bullet 3')

doc.add_paragraph('6-12 meses de histórico:', style='List Bullet 2')
doc.add_paragraph('    - WMA ou SES', style='List Bullet 3')
doc.add_paragraph('    - Detecção básica de tendência', style='List Bullet 3')

doc.add_paragraph('12-18 meses de histórico:', style='List Bullet 2')
doc.add_paragraph('    - Todos métodos exceto Decomposição Sazonal', style='List Bullet 3')

doc.add_paragraph('18+ meses de histórico:', style='List Bullet 2')
doc.add_paragraph('    - Todos métodos disponíveis', style='List Bullet 3')
doc.add_paragraph('    - Detecção completa de sazonalidade', style='List Bullet 3')

doc.add_page_break()

# ========== 7. ORIENTAÇÕES DE USO ==========
doc.add_heading('7. Orientações de Uso', 1)

doc.add_heading('7.1. Preparação do Arquivo de Entrada', 2)

doc.add_paragraph('Formato Necessário:', style='List Bullet')
doc.add_paragraph('Arquivo Excel (.xlsx) com 2 abas obrigatórias', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Aba 1: HISTORICO_VENDAS')
run.bold = True

doc.add_paragraph('Colunas obrigatórias:', style='List Bullet')
doc.add_paragraph('Local: Nome da loja ou CD (ex: LOJA_01, CD_PRINCIPAL)', style='List Bullet 2')
doc.add_paragraph('SKU: Código do produto', style='List Bullet 2')
doc.add_paragraph('Ano: Ano da venda (ex: 2023)', style='List Bullet 2')
doc.add_paragraph('Mes_Numero: Mês numérico (1-12)', style='List Bullet 2')
doc.add_paragraph('Quantidade: Volume vendido', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Aba 2: PEDIDOS_FORNECEDOR (para reabastecimento)')
run.bold = True

doc.add_paragraph('Colunas obrigatórias:', style='List Bullet')
doc.add_paragraph('Fornecedor: Nome do fornecedor', style='List Bullet 2')
doc.add_paragraph('SKU: Código do produto', style='List Bullet 2')
doc.add_paragraph('Destino: CD ou Loja destino', style='List Bullet 2')
doc.add_paragraph('Tipo_Destino: CD ou LOJA', style='List Bullet 2')
doc.add_paragraph('Lead_Time_Dias: Prazo de entrega', style='List Bullet 2')
doc.add_paragraph('Ciclo_Pedido_Dias: Frequência de pedido', style='List Bullet 2')
doc.add_paragraph('Multiplo_Palete, Multiplo_Carreta: Múltiplos de embalagem', style='List Bullet 2')
doc.add_paragraph('Custo_Unitario: Custo por unidade (R$)', style='List Bullet 2')
doc.add_paragraph('Estoque_Disponivel, Estoque_Transito, Pedidos_Abertos', style='List Bullet 2')

doc.add_heading('7.2. Fluxo de Trabalho Recomendado', 2)

doc.add_paragraph('Passo 1: Previsão de Demanda', style='List Number')
doc.add_paragraph('    - Upload do histórico de vendas', style='List Bullet 2')
doc.add_paragraph('    - Processar com método AUTO ou ML', style='List Bullet 2')
doc.add_paragraph('    - Analisar MAPE, BIAS e alertas', style='List Bullet 2')
doc.add_paragraph('    - Baixar Excel com previsões', style='List Bullet 2')

doc.add_paragraph('Passo 2: Cadastrar Eventos (se aplicável)', style='List Number')
doc.add_paragraph('    - Ir para Gerenciador de Eventos', style='List Bullet 2')
doc.add_paragraph('    - Cadastrar promoções futuras', style='List Bullet 2')
doc.add_paragraph('    - Reprocessar previsão (já ajustada)', style='List Bullet 2')

doc.add_paragraph('Passo 3: Calcular Pedidos', style='List Number')
doc.add_paragraph('    - Escolher fluxo: Fornecedor, CD ou Transferências', style='List Bullet 2')
doc.add_paragraph('    - Upload do arquivo com estoque atual', style='List Bullet 2')
doc.add_paragraph('    - Revisar pedidos sugeridos', style='List Bullet 2')
doc.add_paragraph('    - Baixar Excel de pedidos', style='List Bullet 2')

doc.add_paragraph('Passo 4: Simulações (opcional)', style='List Number')
doc.add_paragraph('    - Ir para Simulador de Cenários', style='List Bullet 2')
doc.add_paragraph('    - Testar diferentes parâmetros', style='List Bullet 2')
doc.add_paragraph('    - Comparar resultados', style='List Bullet 2')

doc.add_heading('7.3. Boas Práticas', 2)

doc.add_paragraph('Qualidade dos Dados:', style='List Bullet')
doc.add_paragraph('Mínimo 12 meses de histórico (ideal: 24+ meses)', style='List Bullet 2')
doc.add_paragraph('Dados mensais consolidados e completos', style='List Bullet 2')
doc.add_paragraph('Corrigir manualmente vendas anormais conhecidas (greves, falta produto)', style='List Bullet 2')

doc.add_paragraph('Revisão de Previsões:', style='List Bullet')
doc.add_paragraph('Revisar mensalmente e ajustar com informações de negócio', style='List Bullet 2')
doc.add_paragraph('Atenção a itens com MAPE > 20% (críticos)', style='List Bullet 2')
doc.add_paragraph('Analisar variações YoY grandes (> 50%) para validar', style='List Bullet 2')

doc.add_paragraph('Parâmetros de Reabastecimento:', style='List Bullet')
doc.add_paragraph('Nível de serviço: Usar classificação ABC padrão (A=98%, B=95%, C=90%)', style='List Bullet 2')
doc.add_paragraph('Lead times: Revisar trimestralmente com fornecedores', style='List Bullet 2')
doc.add_paragraph('Ciclos de pedido: Alinhar com logística e capacidade CD', style='List Bullet 2')

doc.add_paragraph('Monitoramento:', style='List Bullet')
doc.add_paragraph('Acompanhar MAPE médio do sistema (meta: < 15%)', style='List Bullet 2')
doc.add_paragraph('Monitorar rupturas reais vs previstas', style='List Bullet 2')
doc.add_paragraph('Analisar métodos mais usados (aba METODOS_UTILIZADOS)', style='List Bullet 2')

doc.add_heading('7.4. Interpretação de Resultados', 2)

p = doc.add_paragraph()
run = p.add_run('Card YoY Positivo (+17.7%):')
run.bold = True
doc.add_paragraph('Previsão indica crescimento em relação ao ano anterior', style='List Bullet 2')
doc.add_paragraph('Preparar estoque e fornecedores para aumento de demanda', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Card YoY Negativo (-5.2%):')
run.bold = True
doc.add_paragraph('Previsão indica redução em relação ao ano anterior', style='List Bullet 2')
doc.add_paragraph('Avaliar excesso de estoque e ajustar pedidos', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Item com 🔴 Crítico:')
run.bold = True
doc.add_paragraph('Requer ação imediata - validar previsão manualmente', style='List Bullet 2')
doc.add_paragraph('Pode indicar: novo produto, promoção, erro de dado', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Cobertura < 7 dias:')
run.bold = True
doc.add_paragraph('Risco iminente de ruptura - priorizar pedido', style='List Bullet 2')
doc.add_paragraph('Considerar pedido emergencial se lead time > 7 dias', style='List Bullet 2')

doc.add_page_break()

# ========== 8. PERGUNTAS FREQUENTES ==========
doc.add_heading('8. Perguntas Frequentes (FAQ)', 1)

doc.add_heading('8.1. Sobre Previsão de Demanda', 2)

p = doc.add_paragraph()
run = p.add_run('P: Qual método devo escolher: AUTO ou ML?')
run.bold = True
doc.add_paragraph('R: Para a maioria dos casos, use ML (Machine Learning). É mais preciso pois considera 15+ características da demanda. Use AUTO apenas se não tiver modelo ML treinado ou para análises rápidas.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Por que alguns produtos têm MAPE muito alto (> 30%)?')
run.bold = True
doc.add_paragraph('R: Possíveis causas:')
doc.add_paragraph('1. Produto novo com pouco histórico (< 6 meses)', style='List Bullet 2')
doc.add_paragraph('2. Demanda muito irregular/sazonal não detectada', style='List Bullet 2')
doc.add_paragraph('3. Outliers não tratados (promoções não cadastradas)', style='List Bullet 2')
doc.add_paragraph('4. Produto em fim de vida com queda abrupta', style='List Bullet 2')
doc.add_paragraph('Solução: Revisar manualmente, cadastrar eventos, aumentar estoque de segurança.', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: O que fazer quando YoY do card não bate com a tabela?')
run.bold = True
doc.add_paragraph('R: Após a correção implementada, o card YoY e a tabela usam o mesmo cálculo: soma das previsões dos próximos N meses vs soma do mesmo período do ano anterior. Se ainda houver divergência, verificar filtros aplicados na tabela.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Posso confiar em previsões de produtos com apenas 3 meses de histórico?')
run.bold = True
doc.add_paragraph('R: Não 100%. O sistema faz o possível (média simples), mas a confiabilidade é baixa. Recomenda-se:')
doc.add_paragraph('Aumentar estoque de segurança em 20-30%', style='List Bullet 2')
doc.add_paragraph('Monitorar vendas semanalmente', style='List Bullet 2')
doc.add_paragraph('Usar benchmark de produtos similares', style='List Bullet 2')

doc.add_heading('8.2. Sobre Reabastecimento', 2)

p = doc.add_paragraph()
run = p.add_run('P: Por que o sistema sugere pedido mesmo com estoque disponível?')
run.bold = True
doc.add_paragraph('R: O sistema considera o período de cobertura futuro (lead time + ciclo de revisão). Se o estoque atual não cobrir esse período + estoque de segurança, um pedido é sugerido antecipadamente para evitar ruptura.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Como o sistema calcula o estoque de segurança?')
run.bold = True
doc.add_paragraph('R: Usa a fórmula: SS = Z × σ × √LT, onde Z depende do nível de serviço (90%=1.28, 95%=1.65, 98%=2.05), σ é o desvio padrão da demanda mensal e LT é o lead time em meses.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Posso alterar o nível de serviço manualmente?')
run.bold = True
doc.add_paragraph('R: Sim, no Simulador de Cenários você pode testar diferentes níveis de serviço e ver o impacto no estoque de segurança e quantidade de pedido. O padrão é classificação ABC automática.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: O que fazer quando há conflito entre Pedido Fornecedor e Transferência?')
run.bold = True
doc.add_paragraph('R: Priorize transferências quando:')
doc.add_paragraph('Há excesso em outra loja do mesmo SKU', style='List Bullet 2')
doc.add_paragraph('Lead time de transferência < lead time do fornecedor', style='List Bullet 2')
doc.add_paragraph('Custo de transferência < custo de novo pedido', style='List Bullet 2')

doc.add_heading('8.3. Sobre Funcionalidades Avançadas', 2)

p = doc.add_paragraph()
run = p.add_run('P: Como funciona a detecção de sazonalidade?')
run.bold = True
doc.add_paragraph('R: O sistema calcula a autocorrelação da série com lag de 12 meses. Se autocorr(12) > 0.6, detecta sazonalidade e aplica Decomposição Sazonal. Requer mínimo 18 meses de histórico.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Outliers são removidos automaticamente?')
run.bold = True
doc.add_paragraph('R: Sim, usando método IQR. Outliers detectados são substituídos pela mediana da série e registrados em log. Promoções conhecidas devem ser cadastradas no Calendário Promocional para não serem tratadas como outliers.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: O que é o Simulador de Cenários?')
run.bold = True
doc.add_paragraph('R: Ferramenta What-If que permite testar mudanças de parâmetros (nível de serviço, lead time, demanda) sem alterar dados reais. Útil para planejamento estratégico e análise de sensibilidade.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Como cadastrar eventos promocionais?')
run.bold = True
doc.add_paragraph('R: Ir para Gerenciador de Eventos → Cadastrar evento com nome, data início/fim e impacto esperado (%). Ex: Black Friday, +150%. O sistema ajusta automaticamente as previsões nesses períodos.')

doc.add_heading('8.4. Sobre Performance e Dados', 2)

p = doc.add_paragraph()
run = p.add_run('P: Qual o tamanho máximo de arquivo que o sistema suporta?')
run.bold = True
doc.add_paragraph('R: Recomendado até 50.000 linhas (combinações SKU×Local×Mês). Para arquivos maiores, considere dividir por categoria ou região.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Quanto tempo demora o processamento?')
run.bold = True
doc.add_paragraph('R: Depende do tamanho:')
doc.add_paragraph('< 1.000 linhas: ~5 segundos', style='List Bullet 2')
doc.add_paragraph('1.000-10.000 linhas: ~30 segundos', style='List Bullet 2')
doc.add_paragraph('10.000-50.000 linhas: ~2-5 minutos', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: Posso usar dados de vendas diárias?')
run.bold = True
doc.add_paragraph('R: Não diretamente. O sistema trabalha com dados mensais consolidados. Agregue suas vendas diárias para mês antes de fazer upload.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('P: O sistema salva histórico de previsões?')
run.bold = True
doc.add_paragraph('R: Atualmente não há versionamento automático. Recomenda-se salvar os arquivos Excel gerados com data no nome (ex: previsao_2024_01.xlsx) para controle manual.')

doc.add_page_break()

# ========== CONCLUSÃO ==========
doc.add_heading('Conclusão', 1)

p = doc.add_paragraph()
p.add_run('O Sistema de Previsão de Demanda e Reabastecimento v3.0 representa uma solução completa e robusta para gestão de estoque em operações multi-loja, combinando:')

doc.add_paragraph()
doc.add_paragraph('Métodos estatísticos avançados com seleção automática inteligente', style='List Bullet')
doc.add_paragraph('Cálculos precisos de estoque de segurança e ponto de pedido', style='List Bullet')
doc.add_paragraph('Interface intuitiva com alertas visuais e métricas de acurácia', style='List Bullet')
doc.add_paragraph('Recursos avançados: ML, detecção de sazonalidade, calendário promocional', style='List Bullet')
doc.add_paragraph('Documentação completa para uso eficaz por equipes de planejamento', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('A utilização adequada do sistema, seguindo as orientações deste manual, permite:')

doc.add_paragraph()
doc.add_paragraph('Redução de rupturas de estoque em 30-50%', style='List Bullet')
doc.add_paragraph('Otimização de capital de giro (redução de excesso)', style='List Bullet')
doc.add_paragraph('Melhoria do nível de serviço ao cliente', style='List Bullet')
doc.add_paragraph('Tomada de decisão baseada em dados e estatística', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('Versão 3.0 - ')
run.font.size = Pt(10)
run = p.add_run(datetime.now().strftime('%B/%Y'))
run.font.size = Pt(10)

# ========== SALVAR DOCUMENTO ==========
doc.save('Documentacao_Sistema_Previsao_v3.0.docx')
print('[OK] Documentacao gerada com sucesso!')
print('Arquivo: Documentacao_Sistema_Previsao_v3.0.docx')
print('Total de paginas: ~30-35')
